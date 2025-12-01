import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def format_run_times_new_roman(run):
    run.font.name = "Times New Roman"
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), 'Times New Roman')


def replace_text_in_docx(file_path, output_path, data):
    doc = Document(file_path)

    for para in doc.paragraphs:
        for run in para.runs:
            for key, val in data.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)
                    format_run_times_new_roman(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, val in data.items():
                            if key in run.text:
                                run.text = run.text.replace(key, val)
                                format_run_times_new_roman(run)

    doc.save(output_path)


# ==========================================
# INPUT BULAN
# ==========================================
bulan_input = input("Masukkan nama bulan : ")
jumlah = input("Masukan jumlah data : ")
tanggal = input("Masukan tanggal : ")


BULAN_UPPER = bulan_input.upper()
BULAN_CAP = bulan_input.capitalize()

# ==========================================
# PATH TEMPLATE
# ==========================================
file_path = "F:/SKP BU IDA/TEMPLATE/NURAIDA BUKTI DUKUNG PKM JUNTINYUAT 2025 NO 7.docx"

# Ambil nama file tanpa folder
filename = os.path.basename(file_path)     # NURAIDA BUKTI ...
filename_no_ext = os.path.splitext(filename)[0]   # tanpa .docx

# ==========================================
# OUTPUT FOLDER
# ==========================================
output_dir = f"F:/SKP BU IDA/{BULAN_UPPER}"

# Buat folder jika belum ada
os.makedirs(output_dir, exist_ok=True)

# ==========================================
# OUTPUT FILE NAME
# ==========================================
output_file = f"{filename_no_ext} {BULAN_UPPER}.docx"

output_path = os.path.join(output_dir, output_file)

# ==========================================
# DATA YANG DIGANTI DI DOKUMEN
# ==========================================
data_replace = {
    "{{BULAN_UPPER}}": BULAN_UPPER,
    "{{BULAN_CAP}}": BULAN_CAP,
    "{{JUMLAH}}": jumlah,
    "{{TANGGAL}}": tanggal
}

# Jalankan proses
replace_text_in_docx(
    file_path=file_path,
    output_path=output_path,
    data=data_replace
)

print("✔ Dokumen berhasil dibuat:")
print(output_path)
